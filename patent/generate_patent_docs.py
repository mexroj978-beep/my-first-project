#!/usr/bin/env python3
"""Patent hujjatlarini yaratish: Annotatsiya (DOCX+PDF) va Dastur kodi (PDF)."""
from __future__ import annotations

import textwrap
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from fpdf import FPDF

OUT = Path(__file__).parent
FONT = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_MONO = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"
TODAY = date.today().strftime("%d.%m.%Y")


ANNOTATION_SECTIONS = [
    (
        "1. Annotatsiya maqsadi",
        "Ushbu hujjat «Maktab Xabarnoma» — maktab turniketi (kirish-chiqish nazorati) "
        "orqali o'quvchi identifikatsiyasi, davomat qayd etish va ota-onalarga Telegram "
        "messenger orqali avtomatik xabar yuborish hamda obuna asosida xizmat ko'rsatish "
        "tizimining ishlash prinsipi haqida qisqacha ma'lumot beradi. Hujjat intellektual "
        "mulk obyektini patentlash uchun taqdim etiladigan annotatsiya (referat) hisoblanadi.",
    ),
    (
        "2. Tizimning umumiy tavsifi",
        "«Maktab Xabarnoma» dasturi maktablar uchun mo'ljallangan dasturiy kompleks bo'lib, "
        "quyidagi komponentlardan iborat: (1) kirish-chiqish turniketi yoki Face ID (yuz "
        "tanish) terminali; (2) markaziy server (FastAPI asosidagi veb-ilova); (3) ma'lumotlar "
        "bazasi (SQLite); (4) Telegram bot; (5) administrator veb-paneli. Tizim o'quvchining "
        "maktabga kirishi yoki chiqishini real vaqt rejimida qayd etadi va bog'langan "
        "ota-onalarga Telegram orqali xabarnoma yuboradi.",
    ),
    (
        "3. Identifikatsiya usuli: Face ID va turniket bog'lanishi",
        "Tizimda o'quvchini tanib olish jarayoni turniket qurilmasi darajasida amalga "
        "oshiriladi. Zamonaviy kirish nazorati terminali ikki usulni qo'llab-quvvatlaydi: "
        "(a) biometrik identifikatsiya — Face ID (yuz tanish) algoritmi yordamida o'quvchi "
        "yuzini kamera orqali skanerlash va bazadagi shablon bilan solishtirish; "
        "(b) RFID/NFC karta orqali identifikatsiya. Face ID terminali muvaffaqiyatli "
        "tanib olgach, qurilma o'quvchiga biriktirilgan noyob identifikatorni (card_id) "
        "markaziy serverga HTTP webhook orqali yuboradi.\n\n"
        "Dasturiy qism biometrik ma'lumotlarni (yuz tasvirlari) saqlamaydi — bu xavfsizlik "
        "va shaxsiy ma'lumotlarni himoya qilish talablariga javob beradi. Server faqat "
        "identifikator qabul qiladi, ma'lumotlar bazasidan o'quvchini topadi va keyingi "
        "jarayonlarni boshlaydi. Shunday qilib, Face ID terminali va Telegram bot quyidagi "
        "zanjir orqali bog'lanadi:\n\n"
        "Face ID terminali → o'quvchini taniydi → card_id + device_code yuboradi → "
        "Server o'quvchini topadi → davomat yozuvi yaratiladi → obuna tekshiriladi → "
        "Telegram bot ota-onaga xabar yuboradi.",
    ),
    (
        "4. Davomat qayd etish jarayoni",
        "Turniket voqeasi POST /webhooks/turnstile manziliga quyidagi ma'lumotlar bilan "
        "keladi: card_id (o'quvchi identifikatori), device_code (qurilma kodi), direction "
        "(kirish/chiqish yo'nalishi), event_time (voqea vaqti). Server X-Webhook-Secret "
        "kaliti orqali so'rovni autentifikatsiya qiladi. AttendanceService moduli "
        "o'quvchini card_id bo'yicha qidiradi, turniket qurilmasini device_code bo'yicha "
        "aniqlaydi, AttendanceEvent jadvaliga yangi yozuv qo'shadi. Keyin o'quvchiga "
        "bog'langan barcha faol ota-onalar ro'yxati olinadi.",
    ),
    (
        "5. Telegram bot bilan integratsiya",
        "Ota-ona Telegram messenjerida maxsus botga ro'yxatdan o'tadi (/register <id> "
        "buyrug'i orqali). Ro'yxatdan o'tishda ota-onaning telegram_chat_id ma'lumotlar "
        "bazasiga yoziladi va bepul sinov muddati boshlanadi (standart: 3 kun). Davomat "
        "voqeasi sodir bo'lganda TelegramService moduli har bir mos ota-onaga HTML formatdagi "
        "xabar yuboradi. Xabarda quyidagilar ko'rsatiladi: o'quvchi ismi, sinfi, maktab "
        "nomi, voqea vaqti, joylashuv (turniket manzili), kirish yoki chiqish holati.\n\n"
        "Bot quyidagi buyruqlarni qo'llab-quvvatlaydi: /start, /register, /status, /pay, "
        "/mystudents, /help. Bot va API server parallel jarayonlarda ishlaydi va bir xil "
        "ma'lumotlar bazasidan foydalanadi.",
    ),
    (
        "6. Obuna va to'lov tizimi",
        "Xabarnoma yuborish obuna holatiga bog'liq: bepul sinov muddati, faol obuna yoki "
        "muddati tugagan holat. SubscriptionService moduli har bir ota-ona uchun holatni "
        "hisoblaydi. To'lov Click.uz integratsiyasi, bank kartasiga o'tkazma yoki demo rejim "
        "orqali amalga oshiriladi. To'lov tasdiqlanganda obuna avtomatik uzaytiriladi va "
        "xabarnomalar qayta yoqiladi.",
    ),
    (
        "7. Administrator paneli",
        "Veb-admin panel (http://server:8000/admin) maktab ma'muriga maktablar, o'quvchilar, "
        "ota-onalar, turniketlar boshqaruvi, obuna sozlamalari va turniket simulyatori "
        "imkonini beradi. Panel REST API orqali ma'lumotlar bazasi bilan ishlaydi.",
    ),
    (
        "8. Texnik arxitektura",
        "Dastur Python 3.12 tilida yozilgan. Backend: FastAPI + SQLAlchemy 2.0 (async) + "
        "SQLite. Telegram integratsiyasi: python-telegram-bot kutubxonasi. Jarayonlar: "
        "multiprocessing orqali API server (Uvicorn, port 8000) va Telegram bot parallel "
        "ishga tushiriladi. Ma'lumotlar modellari: School, Student, Parent, Device, "
        "AttendanceEvent, PaymentOrder, AppSettings.",
    ),
    (
        "9. Ixtironing yangiligi va amaliy ahamiyati",
        "Tizimning o'ziga xos jihati — Face ID/turniket identifikatsiyasi, markaziy server, "
        "obuna sharti asosidagi xabarnoma filtri va Telegram bot orqali ota-onalarga "
        "real vaqtli xabar yetkazishning yagona dasturiy kompleksda birlashtirilishi. "
        "Bir o'quvchiga bir nechta ota-ona bog'lanishi mumkin (ko'p-ota-ona modeli). "
        "Turniket yo'nalishi (kirish/chiqish) qurilma konfiguratsiyasi orqali boshqariladi. "
        "Tizim maktablar, litseylar va boshqa ta'lim muassasalarida o'quvchilar xavfsizligi "
        "va ota-onalar bilan tezkor aloqa uchun qo'llanishi mumkin.",
    ),
]

CODE_FILES = [
    ("start.py — Dasturni ishga tushirish", "start.py"),
    ("app/main.py — FastAPI ilovasi", "app/main.py"),
    ("app/config.py — Sozlamalar", "app/config.py"),
    ("app/database.py — Ma'lumotlar bazasi", "app/database.py"),
    ("app/models/student.py — O'quvchi modeli", "app/models/student.py"),
    ("app/models/parent.py — Ota-ona modeli", "app/models/parent.py"),
    ("app/models/device.py — Turniket modeli", "app/models/device.py"),
    ("app/models/attendance.py — Davomat modeli", "app/models/attendance.py"),
    ("app/models/settings.py — Tizim sozlamalari", "app/models/settings.py"),
    ("app/models/payment.py — To'lov buyurtmasi", "app/models/payment.py"),
    ("app/schemas/__init__.py — API sxemalari (qism)", "app/schemas/__init__.py"),
    ("app/api/webhooks.py — Turniket webhook", "app/api/webhooks.py"),
    ("app/services/attendance.py — Davomat servisi", "app/services/attendance.py"),
    ("app/services/telegram.py — Telegram xabar servisi", "app/services/telegram.py"),
    ("app/services/subscription.py — Obuna servisi", "app/services/subscription.py"),
    ("app/services/payment.py — To'lov servisi", "app/services/payment.py"),
    ("app/bot/handlers.py — Telegram bot buyruqlari", "app/bot/handlers.py"),
    ("app/api/payments.py — To'lov sahifalari", "app/api/payments.py"),
    ("app/api/admin.py — Admin API (qism)", "app/api/admin.py"),
]


def add_annotation_docx() -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "ANNOTATSIYA (REFERAT)\n\n"
        "«Maktab Xabarnoma» — turniket/Face ID identifikatsiyasi va "
        "Telegram bot orqali ota-onalarga avtomatik xabarnoma berish tizimi"
    )
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"\nDasturiy mahsulot nomi: Maktab Xabarnoma v3.0\nSana: {TODAY}").font.size = Pt(11)

    doc.add_paragraph()

    for heading, body in ANNOTATION_SECTIONS:
        h = doc.add_paragraph()
        h.add_run(heading).bold = True
        h.runs[0].font.size = Pt(12)
        for para in body.split("\n\n"):
            p = doc.add_paragraph(para)
            p.paragraph_format.first_line_indent = Cm(1)
            for r in p.runs:
                r.font.size = Pt(12)

    doc.add_paragraph()
    end = doc.add_paragraph("—" * 40)
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Ushbu annotatsiya dasturiy ta'minotning ishlash prinsipi va komponentlararo "
        "o'zaro bog'lanishini qisqacha yoritadi. To'liq dastur kodi alohida PDF hujjatda keltirilgan."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    path = OUT / "Annotatsiya_Maktab_Xabarnoma.docx"
    doc.save(path)
    return path


class PatentPDF(FPDF):
    def __init__(self, title: str):
        super().__init__()
        self.doc_title = title
        self.add_font("DejaVu", "", FONT)
        self.add_font("DejaVu", "B", FONT)
        self.add_font("Mono", "", FONT_MONO)
        self.set_auto_page_break(auto=True, margin=15)

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font("DejaVu", "", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, self.doc_title, align="C", new_x="RIGHT", new_y="TOP")
        self.ln(4)

    def footer(self):
        self.set_y(-12)
        self.set_font("DejaVu", "", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, f"Sahifa {self.page_no()}", align="C")

    def section_title(self, text: str):
        self.ln(4)
        self.set_font("DejaVu", "B", 12)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 7, text)
        self.ln(2)

    def body_text(self, text: str, size: int = 11):
        self.set_font("DejaVu", "", size)
        self.set_text_color(0, 0, 0)
        self.multi_cell(0, 6, text)
        self.ln(2)

    def code_block(self, code: str, filename: str = ""):
        if filename:
            self.set_font("DejaVu", "B", 9)
            self.set_text_color(0, 80, 160)
            self.cell(0, 6, filename, new_x="LMARGIN", new_y="NEXT")
        self.set_font("Mono", "", 7)
        self.set_fill_color(245, 245, 245)
        for line in code.splitlines():
            safe = line.replace("\t", "    ")
            if len(safe) > 110:
                for chunk in textwrap.wrap(safe, 110, break_long_words=False, break_on_hyphens=False):
                    self.cell(0, 4.2, "  " + chunk, fill=True, new_x="LMARGIN", new_y="NEXT")
            else:
                self.cell(0, 4.2, "  " + safe, fill=True, new_x="LMARGIN", new_y="NEXT")
        self.ln(3)


def add_annotation_pdf() -> Path:
    pdf = PatentPDF("Annotatsiya — Maktab Xabarnoma")
    pdf.add_page()
    pdf.set_font("DejaVu", "B", 16)
    pdf.multi_cell(0, 10, "ANNOTATSIYA (REFERAT)", align="C")
    pdf.ln(4)
    pdf.set_font("DejaVu", "B", 12)
    pdf.multi_cell(
        0, 7,
        "«Maktab Xabarnoma» — turniket/Face ID identifikatsiyasi va "
        "Telegram bot orqali ota-onalarga avtomatik xabarnoma berish tizimi",
        align="C",
    )
    pdf.ln(6)
    pdf.set_font("DejaVu", "", 10)
    pdf.cell(0, 6, f"Dasturiy mahsulot: Maktab Xabarnoma v3.0    |    Sana: {TODAY}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    for heading, body in ANNOTATION_SECTIONS:
        pdf.section_title(heading)
        pdf.body_text(body)

    path = OUT / "Annotatsiya_Maktab_Xabarnoma.pdf"
    pdf.output(path)
    return path


def add_code_pdf() -> Path:
    root = OUT.parent
    pdf = PatentPDF("Dastur kodi — Maktab Xabarnoma (asosiy qismlar)")
    pdf.add_page()

    pdf.set_font("DejaVu", "B", 16)
    pdf.multi_cell(0, 10, "DASTUR KODI", align="C")
    pdf.ln(2)
    pdf.set_font("DejaVu", "B", 12)
    pdf.multi_cell(0, 7, "«Maktab Xabarnoma» dasturining asosiy funksiyalarini bajaruvchi modullar", align="C")
    pdf.ln(4)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(
        0, 6,
        f"Sana: {TODAY}\n"
        "Dasturlash tili: Python 3.12\n"
        "Framework: FastAPI, SQLAlchemy, python-telegram-bot\n\n"
        "Ushbu hujjat dasturning to'liq kod bazasining bir qismini — asosiy biznes-logika, "
        "modellar, API, servislar va Telegram bot handlerlarini o'z ichiga oladi.",
    )
    pdf.ln(4)

    pdf.section_title("MUNDARIJA")
    for i, (title, _) in enumerate(CODE_FILES, 1):
        pdf.set_font("DejaVu", "", 10)
        pdf.cell(0, 5, f"{i}. {title}", new_x="LMARGIN", new_y="NEXT")

    pdf.add_page()
    pdf.section_title("TIZIM ARXITEKTURASI (BLOK-SXEMA)")
    pdf.body_text(
        "[Turniket / Face ID terminali]\n"
        "        |  HTTP POST (card_id, device_code, direction)\n"
        "        v\n"
        "[FastAPI Server — /webhooks/turnstile]\n"
        "        |\n"
        "        v\n"
        "[AttendanceService.process()]\n"
        "   1. O'quvchini card_id bo'yicha topish\n"
        "   2. Turniketni device_code bo'yicha topish\n"
        "   3. AttendanceEvent yozuvi yaratish\n"
        "   4. Ota-onalarni StudentParent orqali topish\n"
        "   5. SubscriptionService — obuna tekshiruvi\n"
        "   6. TelegramService — xabar yuborish\n"
        "        |\n"
        "        v\n"
        "[Telegram Bot — ota-ona telefonida xabar]\n\n"
        "Parallel jarayon: Telegram Bot (register, pay, status) <-> SQLite DB <-> Admin Panel"
    )

    for title, rel_path in CODE_FILES:
        fpath = root / rel_path
        if not fpath.exists():
            continue
        code = fpath.read_text(encoding="utf-8")
        pdf.add_page()
        pdf.section_title(title)
        pdf.code_block(code, rel_path)

    # Admin API qolgan qismi alohida (fayl katta)
    admin_path = root / "app/api/admin.py"
    if admin_path.exists():
        lines = admin_path.read_text(encoding="utf-8").splitlines()
        part2 = "\n".join(lines[100:])
        if part2.strip():
            pdf.add_page()
            pdf.section_title("app/api/admin.py — Admin API (davomi)")
            pdf.code_block(part2, "app/api/admin.py (100-qatordan)")

    path = OUT / "Dastur_kodi_Maktab_Xabarnoma.pdf"
    pdf.output(path)
    return path


def main():
    OUT.mkdir(exist_ok=True)
    docx = add_annotation_docx()
    pdf1 = add_annotation_pdf()
    pdf2 = add_code_pdf()
    print(f"Yaratildi:\n  {docx}\n  {pdf1}\n  {pdf2}")
    # Sahifalar soni
    from fpdf import FPDF as F
    # rough page count from file - fpdf doesn't expose easily, use pdfinfo if available
    import subprocess
    for p in [pdf1, pdf2]:
        try:
            r = subprocess.run(["pdfinfo", str(p)], capture_output=True, text=True)
            for line in r.stdout.splitlines():
                if line.startswith("Pages:"):
                    print(f"  {p.name}: {line}")
        except Exception:
            pass


if __name__ == "__main__":
    main()
